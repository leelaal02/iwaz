"""docx_postprocess.py 테스트: 표 행 나눔 금지(cantSplit) 해제."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_postprocess import allow_rows_to_break


def _set_cant_split(row):
    """행에 나눔 금지(w:cantSplit)를 건다(양식 템플릿이 해 두는 설정 재현)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _has_cant_split(row) -> bool:
    trPr = row._tr.find(qn("w:trPr"))
    return trPr is not None and trPr.find(qn("w:cantSplit")) is not None


def test_removes_cant_split_from_all_rows(tmp_path):
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    for row in table.rows:
        _set_cant_split(row)
    assert all(_has_cant_split(r) for r in table.rows)

    removed = allow_rows_to_break(doc)

    assert removed == 3
    assert not any(_has_cant_split(r) for r in table.rows)


def test_rows_without_cant_split_untouched(tmp_path):
    # 나눔 금지가 없던 행은 그대로(제거할 것 없음 → 0 반환)
    doc = Document()
    doc.add_table(rows=2, cols=2)
    assert allow_rows_to_break(doc) == 0


def test_nested_table_rows_also_fixed(tmp_path):
    # 셀 안 중첩표의 행도 재귀로 처리된다
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    _set_cant_split(outer.rows[0])
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    _set_cant_split(nested.rows[0])

    removed = allow_rows_to_break(doc)

    assert removed == 2
    assert not _has_cant_split(outer.rows[0])
    assert not _has_cant_split(nested.rows[0])


def test_round_trip_persists(tmp_path):
    # 저장 후 다시 열어도 cantSplit가 없어야 한다
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _set_cant_split(table.rows[0])
    allow_rows_to_break(doc)
    p = tmp_path / "out.docx"
    doc.save(str(p))
    reopened = Document(str(p))
    assert not _has_cant_split(reopened.tables[0].rows[0])
