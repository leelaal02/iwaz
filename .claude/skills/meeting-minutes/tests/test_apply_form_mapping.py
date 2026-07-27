"""apply_form_mapping.py 테스트: 토큰 블록 라이브러리 + 삽입 + 엔드투엔드."""
import json
from pathlib import Path

import pytest
from docx import Document

from apply_form_mapping import apply_mapping, block_lines, inline_tokens

SAMPLE = Path(__file__).resolve().parent.parent / "examples" / "sample_minutes.json"


def _sample():
    return json.loads(SAMPLE.read_text(encoding="utf-8"))


def _make_form(path, rows=4, cols=2):
    """빈 표 서식(라벨 칸 포함)을 만든다."""
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    table.cell(1, 0).text = "제 목 :"
    table.cell(2, 0).text = "일 시 :"
    doc.save(str(path))


def _cell_texts(out_path, table=0, row=0, col=0):
    t = Document(str(out_path)).tables[table]
    return [p.text for p in t.rows[row].cells[col].paragraphs]


# --- 토큰 블록 라이브러리 단위 테스트 ---------------------------------------
def test_inline_tokens_join():
    # 자동매핑 스칼라는 RichText 슬롯이므로 {{r *_rt }}로 생성한다.
    assert inline_tokens(["title"]) == "{{r title_rt }}"
    assert inline_tokens(["date"]) == "{{r date_rt }}"
    assert inline_tokens(["attendees"]) == "{{r attendees_rt }}"


def test_inline_rejects_list_field():
    with pytest.raises(ValueError, match="목록형"):
        inline_tokens(["discussion"])


def test_block_single_field_no_section_label():
    lines = block_lines(["decisions"])
    assert lines == ["{%p for x in decisions %}", " - {{ x }}", "{%p endfor %}"]
    assert not any(line.startswith("[결정 사항]") for line in lines)


def test_block_single_scalar_is_value_token():
    assert block_lines(["purpose"]) == ["{{r purpose_rt }}"]


def test_block_multi_field_adds_section_labels():
    lines = block_lines(["purpose", "next_meeting"])
    assert lines == ["[회의 목적] {{r purpose_rt }}", "[다음 회의] {{r next_meeting_rt }}"]


def test_block_discussion_uses_numbered_bold_topic():
    lines = block_lines(["discussion"])
    assert lines[0] == "{%p for d in discussion_rt %}"
    assert "{{r d.topic_rt }}" in lines
    assert " - {{ p }}" in lines


def test_block_action_items_owner_due_are_richtext_slots():
    lines = block_lines(["action_items"])
    assert lines[0] == "{%p for a in action_items_rt %}"
    body = " - {{ a.task }} (담당: {{r a.owner_rt }} / 기한: {{r a.due_rt }})"
    assert body in lines


def test_block_multi_list_field_prepends_label_line():
    lines = block_lines(["decisions", "notes"])
    assert lines[0] == "[결정 사항]"
    assert "{%p for x in decisions %}" in lines
    assert "[기타·특이사항]" in lines


def test_unknown_field_raises():
    with pytest.raises(ValueError, match="알 수 없는 field"):
        block_lines(["bogus"])


# --- todo / literal 모드 -----------------------------------------------------
def test_todo_mode_inserts_richtext_todo_token(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(2,0)="일 시 :", 나머지 빈칸
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},   # 빈칸 → todo만
        {"row": 2, "col": 0, "mode": "todo"},   # 라벨칸 → 라벨 뒤 append
    ]}, str(out))
    assert "{{r todo }}" in _cell_texts(out, row=3, col=1)[0]
    label = _cell_texts(out, row=2, col=0)[0]
    assert "일 시 :" in label and "{{r todo }}" in label


def test_todo_mode_needs_no_fields(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    # fields 없이도 동작(예외 없음)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},
    ]}, str(tmp_path / "out.docx"))


def test_todo_mode_end_to_end_red_bold_placeholder(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 1, "mode": "todo"},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    cell = Document(str(final)).tables[0].rows[3].cells[1]
    runs = [(r.text, r.bold, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None)
            for p in cell.paragraphs for r in p.runs if r.text]
    assert any(t == "입력필요" and b and c == "FF0000" for t, b, c in runs)


def test_literal_mode_inserts_given_text(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(3,0) 빈칸
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "literal", "text": "서울 본사 3층"},
    ]}, str(out))
    assert "서울 본사 3층" in _cell_texts(out, row=3, col=0)[0]


def test_literal_mode_requires_text_key(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="text"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 3, "col": 0, "mode": "literal"},
        ]}, str(tmp_path / "out.docx"))


# --- row_repeats (행 반복 표) ------------------------------------------------
def _make_action_table_form(path, merge_task=False):
    """헤더 1행 + 빈 데이터행 1행짜리 실행항목 표 양식."""
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    h = t.rows[0].cells
    h[0].text, h[1].text, h[2].text = "할 일", "담당자", "기한"
    if merge_task:
        r = t.rows[1]
        r.cells[0].merge(r.cells[1])  # 담당자 자리까지 병합(gridSpan)
    doc.save(str(path))


def test_row_repeats_builds_tr_structure(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_action_table_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "owner": 1, "due": 2}},
    ]}, str(out))
    t = Document(str(out)).tables[0]
    assert len(t.rows) == 4  # 헤더 + for행 + 데이터행 + endfor행
    cells = [c.text for row in t.rows for c in row.cells]
    assert any("{%tr for a in action_items_rt %}" in x for x in cells)
    assert any("{%tr endfor %}" in x for x in cells)
    assert "{{ a.task }}" in cells
    assert "{{r a.owner_rt }}" in cells
    assert "{{r a.due_rt }}" in cells


def test_row_repeats_end_to_end_columns_and_todo(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_action_table_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "owner": 1, "due": 2}},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    t = Document(str(final)).tables[0]
    assert len(t.rows) == 3  # 헤더 + 항목 2행
    tasks = [t.rows[i].cells[0].text for i in (1, 2)]
    assert "python-docx 렌더러 PoC 작성" in tasks
    assert "샘플 회의 원문 수집" in tasks
    # due=null 항목의 기한 칸 → 빨간 "입력필요"
    for i in (1, 2):
        if t.rows[i].cells[0].text == "샘플 회의 원문 수집":
            due = t.rows[i].cells[2]
            runs = [(r.text, r.bold, str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None)
                    for p in due.paragraphs for r in p.runs if r.text]
            assert any(tx == "입력필요" and b and c == "FF0000" for tx, b, c in runs)


def test_row_repeats_preserves_gridspan(tmp_path):
    from render_docx_template import render_template
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_action_table_form(tpl, merge_task=True)  # task가 c0~c1 병합
    apply_mapping(str(tpl), {"table": 0, "row_repeats": [
        {"row": 1, "field": "action_items", "cols": {"task": 0, "due": 2}},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    t = Document(str(final)).tables[0]
    assert len(t.rows) == 3  # 헤더 + 2행 (병합 있어도 정상 반복)
    # 각 데이터행에서 병합으로 인해 distinct 셀은 2개(task 병합 + 기한)
    for i in (1, 2):
        distinct = {id(c._tc) for c in t.rows[i].cells}
        assert len(distinct) == 2


def test_row_repeats_unknown_field_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_action_table_form(tpl)
    with pytest.raises(ValueError, match="row_repeats"):
        apply_mapping(str(tpl), {"table": 0, "row_repeats": [
            {"row": 1, "field": "bogus", "cols": {"task": 0}},
        ]}, str(tmp_path / "out.docx"))


# --- 삽입 테스트 -------------------------------------------------------------
def test_inline_appends_after_label(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "inline", "fields": ["title"]},
    ]}, str(out))
    text = _cell_texts(out, row=1, col=0)[0]
    assert "제 목 :" in text            # 라벨 보존
    assert "{{r title_rt }}" in text    # RichText 토큰이 라벨 뒤에


def test_block_single_fills_cell(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    lines = _cell_texts(out, row=3, col=0)
    assert "{%p for x in decisions %}" in lines
    assert " - {{ x }}" in lines


def test_block_multi_keeps_order_and_labels(tmp_path):
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["purpose", "next_meeting"]},
    ]}, str(out))
    lines = _cell_texts(out, row=3, col=0)
    assert lines[0] == "[회의 목적] {{r purpose_rt }}"
    assert lines[1] == "[다음 회의] {{r next_meeting_rt }}"


def test_block_preserves_labeled_cell(tmp_path):
    """라벨 칸("제 목 :")에 block을 넣어도 라벨을 지우지 않고 그 뒤에 블록을 붙인다."""
    tpl = tmp_path / "form.docx"
    out = tmp_path / "out.docx"
    _make_form(tpl)  # cell(1,0) == "제 목 :"
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    lines = _cell_texts(out, row=1, col=0)
    assert lines[0] == "제 목 :"                       # 라벨 보존(첫 문단 유지)
    assert "{%p for x in decisions %}" in lines        # 블록은 라벨 뒤에
    assert lines.index("{%p for x in decisions %}") > 0


def test_out_of_range_row_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl, rows=4)
    with pytest.raises(IndexError, match="행"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 99, "col": 0, "mode": "inline", "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_out_of_range_table_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(IndexError, match="표 인덱스"):
        apply_mapping(str(tpl), {"table": 5, "fills": [
            {"row": 0, "col": 0, "mode": "inline", "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_missing_fill_key_raises_friendly_error(tmp_path):
    # mode 키가 빠진 fills 항목 → 원시 KeyError 대신 안내 메시지(ValueError)
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="필수 키 'mode'"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 1, "col": 0, "fields": ["title"]},
        ]}, str(tmp_path / "out.docx"))


def test_missing_paragraph_key_raises_friendly_error(tmp_path):
    tpl = tmp_path / "pform.docx"
    _make_para_form(tpl)
    with pytest.raises(ValueError, match="필수 키 'fields'"):
        apply_mapping(str(tpl), {"paragraphs": [
            {"para": 1, "mode": "inline"},
        ]}, str(tmp_path / "out.docx"))


def test_unknown_field_in_apply_raises(tmp_path):
    tpl = tmp_path / "form.docx"
    _make_form(tpl)
    with pytest.raises(ValueError, match="알 수 없는 field"):
        apply_mapping(str(tpl), {"table": 0, "fills": [
            {"row": 3, "col": 0, "mode": "block", "fields": ["nope"]},
        ]}, str(tmp_path / "out.docx"))


# --- 문단 기반 양식 삽입 -----------------------------------------------------
def _make_para_form(path):
    """개요/참석자/회의내용 문단 양식(표 없음)을 만든다."""
    doc = Document()
    doc.add_paragraph("개요")          # 0
    doc.add_paragraph("ㅇ (목적)")     # 1
    doc.add_paragraph("ㅇ (일시)")     # 2
    doc.add_paragraph("회의내용")      # 3
    doc.add_paragraph("ㅇ")            # 4
    doc.add_paragraph("ㅇ")            # 5
    doc.save(str(path))


def _para_texts(out_path):
    return [p.text for p in Document(str(out_path)).paragraphs]


def test_paragraph_inline_appends_after_label(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 1, "mode": "inline", "fields": ["purpose"]},
        {"para": 2, "mode": "inline", "fields": ["date"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert "ㅇ (목적) {{r purpose_rt }}" in texts
    assert "ㅇ (일시) {{r date_rt }}" in texts


def test_paragraph_block_inserts_following_paragraphs(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 4, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    texts = _para_texts(out)
    # block 첫 라인은 대상 문단에, 나머지는 바로 뒤 문단으로
    assert "{%p for x in decisions %}" in texts
    assert " - {{ x }}" in texts
    assert "{%p endfor %}" in texts
    # 회의내용 헤딩은 보존되고, 뒤의 다른 ㅇ 문단도 남아 있다
    assert "회의내용" in texts


def test_paragraph_block_multi_field_labels(tmp_path):
    tpl = tmp_path / "pform.docx"
    out = tmp_path / "out.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 4, "mode": "block", "fields": ["decisions", "notes"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert "[결정 사항]" in texts
    assert "[기타·특이사항]" in texts


def test_block_preserves_labeled_paragraph(tmp_path):
    """라벨 문단에 block을 넣어도 라벨을 지우지 않고 바로 뒤에 블록을 삽입한다."""
    doc = Document()
    doc.add_paragraph("ㅇ 결정사항")   # 0: 라벨 문단
    doc.add_paragraph("뒤 문단")        # 1: 이후 문단 보존 확인용
    tpl = tmp_path / "pform.docx"
    doc.save(str(tpl))
    out = tmp_path / "out.docx"
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 0, "mode": "block", "fields": ["decisions"]},
    ]}, str(out))
    texts = _para_texts(out)
    assert texts[0] == "ㅇ 결정사항"                    # 라벨 보존(그 자리 유지)
    assert texts[1] == "{%p for x in decisions %}"      # 블록이 라벨 바로 뒤
    assert "뒤 문단" in texts                           # 이후 문단도 보존


def test_paragraph_out_of_range_raises(tmp_path):
    tpl = tmp_path / "pform.docx"
    _make_para_form(tpl)
    with pytest.raises(IndexError, match="문단 인덱스"):
        apply_mapping(str(tpl), {"paragraphs": [
            {"para": 999, "mode": "inline", "fields": ["purpose"]},
        ]}, str(tmp_path / "out.docx"))


def test_mixed_table_and_paragraph_fills(tmp_path):
    """표 셀(title) + 문단(purpose)을 한 번에 채운다."""
    doc = Document()
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "회의록"
    doc.add_paragraph("ㅇ (목적)")  # 문단 index 0 (표 안 문단은 제외되므로)
    tpl = tmp_path / "mixed.docx"
    doc.save(str(tpl))
    out = tmp_path / "out.docx"
    apply_mapping(str(tpl), {
        "table": 0,
        "fills": [{"row": 0, "col": 0, "mode": "inline", "fields": ["title"]}],
        "paragraphs": [{"para": 0, "mode": "inline", "fields": ["purpose"]}],
    }, str(out))
    doc_out = Document(str(out))
    assert "회의록 {{r title_rt }}" in doc_out.tables[0].rows[0].cells[0].text
    assert "ㅇ (목적) {{r purpose_rt }}" in [p.text for p in doc_out.paragraphs]


def test_paragraph_block_end_to_end(tmp_path):
    """문단 block 삽입 → render_docx_template로 실제 값이 채워지는지."""
    from render_docx_template import render_template

    tpl = tmp_path / "pform.docx"
    tokenized = tmp_path / "pform_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_para_form(tpl)
    apply_mapping(str(tpl), {"paragraphs": [
        {"para": 1, "mode": "inline", "fields": ["purpose"]},
        {"para": 4, "mode": "block", "fields": ["decisions"]},
        {"para": 5, "mode": "block", "fields": ["action_items"]},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))
    all_text = "\n".join(_para_texts(final))
    assert "docx 변환은 python-docx로 진행" in all_text     # decisions
    assert "python-docx 렌더러 PoC 작성" in all_text         # action_items


# --- 엔드투엔드: apply → render_docx_template ---------------------------------
def test_end_to_end_fills_values(tmp_path):
    from render_docx_template import render_template

    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl, rows=5)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 1, "col": 0, "mode": "inline", "fields": ["title"]},
        {"row": 2, "col": 0, "mode": "inline", "fields": ["date"]},
        {"row": 3, "col": 0, "mode": "block", "fields": ["decisions"]},
        {"row": 4, "col": 0, "mode": "block", "fields": ["action_items"]},
    ]}, str(tokenized))
    render_template(str(tokenized), _sample(), str(final))

    table = Document(str(final)).tables[0]
    all_text = "\n".join(
        p.text for row in table.rows for cell in row.cells for p in cell.paragraphs
    )
    assert "2026 3분기 제품 로드맵 회의" in all_text     # title inline
    assert "2026-07-16 14:00" in all_text               # date inline
    assert "docx 변환은 python-docx로 진행" in all_text  # decisions 반복
    assert "python-docx 렌더러 PoC 작성" in all_text     # action_items 반복


def test_end_to_end_empty_list_leaves_blank(tmp_path):
    from render_docx_template import render_template

    data = _sample()
    data["notes"] = []
    tpl = tmp_path / "form.docx"
    tokenized = tmp_path / "form_tokenized.docx"
    final = tmp_path / "final.docx"
    _make_form(tpl, rows=4)
    apply_mapping(str(tpl), {"table": 0, "fills": [
        {"row": 3, "col": 0, "mode": "block", "fields": ["notes"]},
    ]}, str(tokenized))
    render_template(str(tokenized), data, str(final))

    lines = _cell_texts(final, row=3, col=0)
    # notes가 비면 반복 문단이 생성되지 않아 " - ..." 라인이 없다
    assert not any(line.strip().startswith("-") for line in lines)
