"""inspect_template.py 테스트: 좌표·is_empty·병합 원점 dedup·has_tokens."""
import pytest
from docx import Document

from inspect_template import inspect_template


def _make_form(path, *, tokens=False):
    """토큰 없는(또는 있는) 표 서식 최소 양식을 만든다.

    구조:
      row0: 병합된 제목 셀 "회 의 록"
      row1: "제 목 :" | (빈 값 칸)
      row2: "일 시 :" | "참 가 자"
      row3: (빈칸, tokens=True면 토큰 삽입) | (빈칸)
    """
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "회 의 록"
    table.cell(0, 0).merge(table.cell(0, 1))  # 가로 병합
    table.cell(1, 0).text = "제 목 :"
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "일 시 :"
    table.cell(2, 1).text = "참 가 자"
    if tokens:
        table.cell(3, 0).text = "{{ title }}"
    doc.save(str(path))


def test_no_tokens_detected(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p, tokens=False)
    result = inspect_template(str(p))
    assert result["has_tokens"] is False


def test_tokens_detected(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p, tokens=True)
    result = inspect_template(str(p))
    assert result["has_tokens"] is True


def test_tokens_in_header_detected(tmp_path):
    """머리말에만 토큰이 있어도 탐지한다 — 본문만 훑으면 놓친다."""
    doc = Document()
    doc.add_table(rows=1, cols=1)  # 본문엔 토큰 없음
    header = doc.sections[0].header
    header.paragraphs[0].text = "{{ title }}"
    p = tmp_path / "header_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_tokens_in_footer_detected(tmp_path):
    """꼬리말에만 토큰이 있어도 탐지한다."""
    doc = Document()
    doc.add_paragraph("본문 문단")
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "{% for a in attendees %}{{ a }}{% endfor %}"
    p = tmp_path / "footer_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_tokens_in_nested_table_detected(tmp_path):
    """표 셀 안의 중첩표에 든 토큰도 탐지한다 — cell.text는 중첩표를 안 훑는다."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "{{ purpose }}"
    p = tmp_path / "nested_form.docx"
    doc.save(str(p))
    assert inspect_template(str(p))["has_tokens"] is True


def test_table_dimensions(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    table = inspect_template(str(p))["tables"][0]
    assert table["index"] == 0
    assert table["rows"] == 4
    assert table["cols"] == 2


def test_merged_origin_deduped(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    cells = inspect_template(str(p))["tables"][0]["cells"]
    # row0은 가로 병합 → 원점 (0,0) 하나만, merged=True
    row0 = [c for c in cells if c["row"] == 0]
    assert len(row0) == 1
    assert row0[0]["col"] == 0
    assert row0[0]["merged"] is True
    assert row0[0]["text"] == "회 의 록"


def test_empty_cell_flagged(tmp_path):
    p = tmp_path / "form.docx"
    _make_form(p)
    cells = inspect_template(str(p))["tables"][0]["cells"]
    by_coord = {(c["row"], c["col"]): c for c in cells}
    assert by_coord[(1, 1)]["is_empty"] is True   # 빈 값 칸
    assert by_coord[(1, 0)]["is_empty"] is False  # "제 목 :"
    assert by_coord[(2, 1)]["text"] == "참 가 자"


def test_non_docx_raises(tmp_path):
    bad = tmp_path / "form.txt"
    bad.write_text("not a docx", encoding="utf-8")
    with pytest.raises(ValueError, match="docx"):
        inspect_template(str(bad))


def test_no_tables_returns_empty(tmp_path):
    doc = Document()
    doc.add_paragraph("표 없는 문단전용 양식")
    p = tmp_path / "noform.docx"
    doc.save(str(p))
    result = inspect_template(str(p))
    assert result["tables"] == []
    assert result["has_tokens"] is False


def test_paragraphs_dumped_with_index(tmp_path):
    doc = Document()
    doc.add_paragraph("개요")
    doc.add_paragraph("ㅇ (목적)")
    doc.add_paragraph("")  # 빈 문단도 인덱스 정확성을 위해 포함
    doc.add_paragraph("ㅇ (일시)")
    p = tmp_path / "para_form.docx"
    doc.save(str(p))
    paras = inspect_template(str(p))["paragraphs"]
    by_idx = {x["index"]: x for x in paras}
    assert by_idx[0]["text"] == "개요"
    assert by_idx[1]["text"] == "ㅇ (목적)"
    assert by_idx[2]["is_empty"] is True
    assert by_idx[3]["text"] == "ㅇ (일시)"
