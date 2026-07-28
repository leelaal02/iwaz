"""[4-보강] 매핑 적용: 양식 복사본에 검증된 docxtpl 토큰 블록을 삽입.

자동 매핑 경로의 ③단계. Claude가 만든 mapping.json대로 표의 지정 칸에
**검증된 토큰 텍스트만** 넣는다(Claude는 토큰 문법을 짜지 않는다 → 문법 오류 차단).
삽입 후에는 기존 render_docx_template.py가 무수정으로 값을 채운다.

원본 양식은 절대 수정하지 않는다 — 호출자가 준 복사본 경로(out_path)에만 저장.
"""
import json
import sys
from pathlib import Path

from inspect_template import is_shaded
from normalize_input import resolve_input_path

# --- 토큰 블록 라이브러리 (설계 §4) ------------------------------------------
# 서식(빨강 "입력필요"·소제목 굵게)이 필요한 값은 RichText 슬롯이라 반드시
# `{{r ... }}`(r 접두사)로 넣는다. 일반 `{{ }}`에 RichText를 주면 run이 중첩되어
# 빈칸으로 깨진다(docxtpl 제약). 대응 컨텍스트 키는 render_docx_template.build_context
# 의 `*_rt`/`discussion_rt`/`action_items_rt`. 항상-존재하는 평문(points·task 등)은
# 일반 `{{ }}`로 둔다.
_SCALAR_TOKEN = {
    "title": "{{r title_rt }}",
    "date": "{{r date_rt }}",
    "attendees": "{{r attendees_rt }}",
    "purpose": "{{r purpose_rt }}",
    "next_meeting": "{{r next_meeting_rt }}",
}

# 목록형 field의 block 본문(문단 라인 목록). {%p%}는 문단 단위 반복 태그.
_LIST_BLOCK = {
    "discussion": [
        "{%p for d in discussion_rt %}",
        "{{r d.topic_rt }}",  # "N. 주제" 넘버링+굵게
        "{%p for p in d.points %}",
        " - {{ p }}",
        "{%p endfor %}",
        "{%p endfor %}",
    ],
    "decisions": [
        "{%p for x in decisions %}",
        " - {{ x }}",
        "{%p endfor %}",
    ],
    "action_items": [
        "{%p for a in action_items_rt %}",
        " - {{ a.task }} (담당: {{r a.owner_rt }} / 기한: {{r a.due_rt }})",
        "{%p endfor %}",
    ],
    "notes": [
        "{%p for n in notes %}",
        " - {{ n }}",
        "{%p endfor %}",
    ],
}

# 행 반복 표(row_repeats)용: 리스트형 field → 반복 대상 컨텍스트 키·루프 변수·
# 하위필드별 컬럼 토큰. task는 항상-존재 평문({{ }}), owner/due는 RichText 슬롯({{r }}).
_ROW_REPEAT = {
    "action_items": {
        "iter": "action_items_rt",
        "var": "a",
        "col_tokens": {
            "task": "{{ a.task }}",
            "owner": "{{r a.owner_rt }}",
            "due": "{{r a.due_rt }}",
        },
    },
}

# 한 칸에 여러 field가 들어갈 때 구분용 섹션 라벨(block 복수 field에서만 붙임).
_SECTION_LABEL = {
    "title": "[제목]",
    "date": "[일시]",
    "attendees": "[참석자]",
    "purpose": "[회의 목적]",
    "next_meeting": "[다음 회의]",
    "discussion": "[논의 내용]",
    "decisions": "[결정 사항]",
    "action_items": "[실행 항목]",
    "notes": "[기타·특이사항]",
}

ALLOWED_FIELDS = set(_SECTION_LABEL)  # 고정 9항목 어휘


def _validate_fields(fields) -> None:
    if not fields:
        raise ValueError("fill의 'fields'가 비어 있습니다.")
    for f in fields:
        if f not in ALLOWED_FIELDS:
            raise ValueError(
                f"알 수 없는 field '{f}'. 허용 어휘: "
                + ", ".join(sorted(ALLOWED_FIELDS))
            )


def inline_tokens(fields) -> str:
    """inline mode: 스칼라 field의 값 토큰을 공백으로 이어 한 줄로."""
    _validate_fields(fields)
    parts = []
    for f in fields:
        if f not in _SCALAR_TOKEN:
            raise ValueError(
                f"'{f}'는 목록형이라 inline 모드로 넣을 수 없습니다. block 모드를 쓰세요."
            )
        parts.append(_SCALAR_TOKEN[f])
    return " ".join(parts)


def block_lines(fields) -> list:
    """block mode: field 순서대로 토큰 문단 라인 목록을 만든다.

    field가 하나면 섹션 라벨 생략, 둘 이상이면 각 field에 섹션 라벨 줄을 붙인다.
    """
    _validate_fields(fields)
    single = len(fields) == 1
    lines = []
    for f in fields:
        if f in _SCALAR_TOKEN:  # 스칼라·title·date는 값 토큰 한 줄
            token = _SCALAR_TOKEN[f]
            lines.append(token if single else f"{_SECTION_LABEL[f]} {token}")
        else:  # 목록형: 반복 블록
            if not single:
                lines.append(_SECTION_LABEL[f])
            lines.extend(_LIST_BLOCK[f])
    return lines


# --- docx 조작 --------------------------------------------------------------
def _clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


_TODO_TOKEN = "{{r todo }}"  # build_context의 RichText `todo`("입력필요" 빨강·굵게)


def _append_token(paragraph, token: str) -> None:
    """paragraph 끝에 임의의 토큰/텍스트를 덧붙인다(라벨 텍스트 보존)."""
    sep = " " if paragraph.text.strip() else ""  # 라벨이 있으면 한 칸 띄움
    paragraph.add_run(sep + token)


def _append_inline(paragraph, fields) -> None:
    """paragraph 끝에 inline 토큰을 덧붙인다(라벨 텍스트 보존)."""
    _append_token(paragraph, inline_tokens(fields))


def _insert_paragraph_after(paragraph, text: str):
    """본문에서 paragraph 바로 뒤에 새 문단을 만들어 반환한다.

    python-docx에 공개 API가 없어 XML(addnext)로 삽입한다. 표 셀이 아니라
    문단 기반 양식의 block 채움에서 여러 문단을 순서대로 끼워 넣을 때 쓴다.
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _apply_inline(cell, fields) -> None:
    _append_inline(cell.paragraphs[0], fields)


def _apply_todo(cell) -> None:
    """라벨만 있고 데이터가 없는 칸에 빨간 "입력필요"(RichText) 토큰을 넣는다."""
    _append_token(cell.paragraphs[0], _TODO_TOKEN)


def _apply_literal(cell, text) -> None:
    """원문에서 찾은 값 등 지정 문자열을 평문으로 넣는다."""
    _append_token(cell.paragraphs[0], text)


def _apply_block(cell, fields) -> None:
    """표 셀에 block 토큰을 삽입한다.

    라벨이 있는 칸("결정 사항:" 등)은 지우지 않고 그 뒤에 블록을 새 문단으로
    넣는다(inline과 동일한 라벨 보존 규칙). 빈 칸이면 첫 라인을 기존 첫 문단에
    넣어 셀 서식을 유지한다.
    """
    lines = block_lines(fields)
    p0 = cell.paragraphs[0]
    if p0.text.strip():  # 라벨 보존: 기존 텍스트 유지, 블록은 전부 뒤에
        for line in lines:
            cell.add_paragraph(line)
    else:  # 빈 칸: 첫 라인은 기존 첫 문단에(서식 유지), 나머지는 뒤에
        _clear_runs(p0)
        p0.add_run(lines[0])
        for line in lines[1:]:
            cell.add_paragraph(line)


def _apply_block_paragraph(paragraph, fields) -> None:
    """문단 기반 양식에 block 토큰을 삽입한다.

    라벨이 있는 문단("ㅇ 논의내용" 등)은 지우지 않고 그 뒤에 블록을 새 문단으로
    넣는다. 빈 문단이면 첫 라인을 그 문단에 넣고 나머지를 뒤에 삽입한다.
    """
    lines = block_lines(fields)
    if paragraph.text.strip():  # 라벨 보존: 기존 문단 유지, 블록은 전부 뒤에
        prev = paragraph
        for line in lines:
            prev = _insert_paragraph_after(prev, line)
    else:  # 빈 문단: 첫 라인은 그 문단에, 나머지는 뒤에
        _clear_runs(paragraph)
        paragraph.add_run(lines[0])
        prev = paragraph
        for line in lines[1:]:
            prev = _insert_paragraph_after(prev, line)


def _check_not_shaded(cell, entry: dict, where: str) -> None:
    """배경색이 칠해진 칸에 값을 넣으려 하면 막는다(색 무관 — 회색뿐 아니라 어떤 색이든).

    색칠된 칸은 거의 항상 라벨/헤더이고, 값은 인접한 색 없는 칸이나 다음 행·문단에
    들어간다. block/inline은 라벨을 지우지 않고 뒤에 이어 붙이므로 색칠된 칸에
    매핑해도 오류 없이 통과해 **라벨 칸 안에 본문이 박히는** 결과가 나온다.
    이 조용한 오배치를 실행 시점에 시끄럽게 만든다.

    매핑 항목에 `"allow_shaded": true`가 있으면 통과시키되, 이 우회는
    **사용자가 그 칸에 넣으라고 명시적으로 요청했을 때만** 쓴다. 넣을 자리를
    못 찾았다는 이유로 붙이는 용도가 아니다(그 경우 다른 빈 자리로 옮긴다).
    """
    if entry.get("allow_shaded"):
        return
    if is_shaded(cell._tc):
        raise ValueError(
            f"{where}: 배경색이 칠해진 칸에 값을 넣으려 합니다. 색칠된 칸은 "
            "라벨/헤더이므로 값은 같은 행의 색 없는 칸이나 바로 다음 빈 행·문단에 "
            "매핑하세요(구조 JSON의 shaded=false 자리). 정말 이 칸에 넣어야 하면 "
            '해당 매핑 항목에 "allow_shaded": true를 추가하세요.'
        )


def _require(entry: dict, key: str, where: str):
    """매핑 항목에서 필수 키를 꺼낸다. 없으면 원시 KeyError 대신 안내 메시지."""
    if key not in entry:
        raise ValueError(
            f"매핑 항목({where})에 필수 키 '{key}'가 없습니다. "
            "각 fills 항목은 row·col·mode·fields를, paragraphs 항목은 para·mode·fields를 가져야 합니다."
        )
    return entry[key]


def _apply_table_fills(doc, mapping) -> None:
    """표 셀 채움(`fills`). 표가 없으면 fills가 있을 때만 오류."""
    fills = mapping.get("fills", [])
    if not fills:
        return
    ti = mapping.get("table", 0)
    if ti < 0 or ti >= len(doc.tables):
        raise IndexError(
            f"표 인덱스 {ti}가 범위를 벗어났습니다(문서의 표는 {len(doc.tables)}개)."
        )
    table = doc.tables[ti]
    n_rows = len(table.rows)
    for fill in fills:
        r = _require(fill, "row", "fills")
        c = _require(fill, "col", "fills")
        mode = _require(fill, "mode", "fills")
        if r < 0 or r >= n_rows:
            raise IndexError(
                f"행 {r}이(가) 표 범위를 벗어났습니다(표 {ti}의 행은 {n_rows}개)."
            )
        row_cells = table.rows[r].cells
        if c < 0 or c >= len(row_cells):
            raise IndexError(
                f"열 {c}이(가) 표 범위를 벗어났습니다(행 {r}의 열은 {len(row_cells)}개)."
            )
        cell = row_cells[c]
        _check_not_shaded(cell, fill, f"fills[{r},{c}]")
        if mode == "inline":
            _apply_inline(cell, _require(fill, "fields", "fills"))
        elif mode == "block":
            _apply_block(cell, _require(fill, "fields", "fills"))
        elif mode == "todo":
            _apply_todo(cell)
        elif mode == "literal":
            _apply_literal(cell, _require(fill, "text", "fills"))
        else:
            raise ValueError(
                f"알 수 없는 mode '{mode}' (inline|block|todo|literal 중 하나)."
            )


def _apply_paragraph_fills(doc, mapping) -> None:
    """문단 채움(`paragraphs`). 인덱스를 먼저 스냅샷해 삽입에 따른 인덱스 밀림을 피한다."""
    para_fills = mapping.get("paragraphs", [])
    if not para_fills:
        return
    paras = doc.paragraphs
    n = len(paras)
    # block 삽입이 뒤 인덱스를 밀기 전에 대상 문단 객체를 먼저 확보한다.
    targets = []
    for pf in para_fills:
        idx = _require(pf, "para", "paragraphs")
        if idx < 0 or idx >= n:
            raise IndexError(
                f"문단 인덱스 {idx}가 범위를 벗어났습니다(본문 문단은 {n}개)."
            )
        targets.append((paras[idx], pf))
    for paragraph, pf in targets:
        mode = _require(pf, "mode", "paragraphs")
        if mode == "inline":
            _append_inline(paragraph, _require(pf, "fields", "paragraphs"))
        elif mode == "block":
            _apply_block_paragraph(paragraph, _require(pf, "fields", "paragraphs"))
        elif mode == "todo":
            _append_token(paragraph, _TODO_TOKEN)
        elif mode == "literal":
            _append_token(paragraph, _require(pf, "text", "paragraphs"))
        else:
            raise ValueError(
                f"알 수 없는 mode '{mode}' (inline|block|todo|literal 중 하나)."
            )


def _make_tag_row(data_tr, tag: str):
    """데이터 행 XML을 복제해 모든 셀을 비우고 첫 셀에 `{%tr ...%}` 태그만 넣는다.

    셀 수·gridSpan을 데이터 행과 동일하게 유지해야 표가 깨지지 않으므로 deepcopy를 쓴다.
    """
    import copy

    from docx.oxml.ns import qn

    new = copy.deepcopy(data_tr)
    for tc in new.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
    first_p = new.findall(qn("w:tc"))[0].find(qn("w:p"))
    run = first_p.makeelement(qn("w:r"), {})
    wt = first_p.makeelement(qn("w:t"), {})
    wt.text = tag
    run.append(wt)
    first_p.append(run)
    return new


def _set_cell_token(cell, token: str) -> None:
    """표 셀의 첫 문단을 비우고 토큰 하나만 넣는다(데이터 행 컬럼 채움용)."""
    p = cell.paragraphs[0]
    _clear_runs(p)
    p.add_run(token)


def _apply_row_repeats(doc, mapping) -> None:
    """`row_repeats`: 리스트형 field를 `{%tr%}` 3행 구조로 표에 펼친다(컬럼별 채움).

    데이터 행의 지정 컬럼에 하위필드 토큰을 넣고, 그 앞/뒤에 for·endfor 태그 행을
    삽입한다. 렌더 시 docxtpl가 데이터 행을 항목 수만큼 반복하고 태그 행은 삭제한다.
    """
    repeats = mapping.get("row_repeats", [])
    if not repeats:
        return
    default_ti = mapping.get("table", 0)
    for entry in repeats:
        field = _require(entry, "field", "row_repeats")
        if field not in _ROW_REPEAT:
            raise ValueError(
                f"row_repeats의 field '{field}'는 행 반복 대상이 아닙니다. "
                "허용: " + ", ".join(sorted(_ROW_REPEAT))
            )
        spec = _ROW_REPEAT[field]
        ti = entry.get("table", default_ti)
        if ti < 0 or ti >= len(doc.tables):
            raise IndexError(
                f"표 인덱스 {ti}가 범위를 벗어났습니다(문서의 표는 {len(doc.tables)}개)."
            )
        table = doc.tables[ti]
        row = _require(entry, "row", "row_repeats")
        if row < 0 or row >= len(table.rows):
            raise IndexError(
                f"행 {row}이(가) 표 범위를 벗어났습니다(표 {ti}의 행은 {len(table.rows)}개)."
            )
        cols = _require(entry, "cols", "row_repeats")
        data_row = table.rows[row]
        row_cells = data_row.cells
        for subfield, col_idx in cols.items():
            if subfield not in spec["col_tokens"]:
                raise ValueError(
                    f"row_repeats field '{field}'에 없는 하위필드 '{subfield}'. "
                    "허용: " + ", ".join(sorted(spec["col_tokens"]))
                )
            if col_idx < 0 or col_idx >= len(row_cells):
                raise IndexError(
                    f"열 {col_idx}이(가) 표 범위를 벗어났습니다(행 {row}의 열은 {len(row_cells)}개)."
                )
            _check_not_shaded(
                row_cells[col_idx], entry, f"row_repeats[{row},{col_idx}]"
            )
            _set_cell_token(row_cells[col_idx], spec["col_tokens"][subfield])
        data_tr = data_row._tr
        for_tag = f"{{%tr for {spec['var']} in {spec['iter']} %}}"
        data_tr.addprevious(_make_tag_row(data_tr, for_tag))
        data_tr.addnext(_make_tag_row(data_tr, "{%tr endfor %}"))


def apply_mapping(template_path: str, mapping: dict, out_path: str) -> None:
    """mapping대로 양식 복사본에 토큰을 삽입해 out_path에 저장한다.

    `fills`(표 셀)·`paragraphs`(본문 문단)·`row_repeats`(행 반복 표) 중
    있는 것만 적용한다.
    """
    from docx import Document

    resolved = resolve_input_path(template_path, must_exist=True)
    doc = Document(str(resolved))
    _apply_table_fills(doc, mapping)
    _apply_paragraph_fills(doc, mapping)
    _apply_row_repeats(doc, mapping)
    doc.save(out_path)


def main() -> None:
    template_path, mapping_path, out_path = sys.argv[1], sys.argv[2], sys.argv[3]
    mapping = json.loads(Path(mapping_path).read_text(encoding="utf-8"))
    apply_mapping(template_path, mapping, out_path)
    print(f"토큰 삽입 완료(토큰화 양식): {out_path}")


if __name__ == "__main__":
    main()
