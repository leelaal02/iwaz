"""[4-보강] 양식 구조 덤프: 토큰 없는 표 서식(.docx) → 구조 JSON.

자동 매핑 경로의 ①단계. 표·칸 라벨·좌표·빈칸·병합·토큰 유무를 JSON으로
출력해 Claude가 9항목을 어느 칸에 넣을지 판단(mapping.json)하는 입력으로 쓴다.
스크립트는 순수 기계적(구조만 보고), 의미 판단은 하지 않는다.

병합 셀은 원점(top-left) 1회만 출력하고, 병합 중복 셀은 제외한다.
표만 처리한다(v1). 표 밖 문단은 매핑 대상이 아니다.
"""
import json
import sys

from normalize_input import resolve_input_path


def _iter_container_texts(container):
    """컨테이너(본문·머리말/꼬리말·표 셀)의 문단 텍스트를 중첩표까지 재귀로 낸다.

    Document 본문·`_Header`/`_Footer`·`_Cell`은 모두 `.paragraphs`·`.tables`를
    공유하므로 한 함수로 처리한다. paragraph.text는 런을 이어붙이므로 한 문단 안의
    토큰은 쪼개지지 않는다. 셀 안의 중첩표까지 재귀해 어디에 있든 토큰을 놓치지 않는다.
    """
    for p in container.paragraphs:
        yield p.text
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_container_texts(cell)


def _iter_all_texts(doc):
    """문서 본문 + 모든 섹션의 머리말/꼬리말 텍스트를 재귀로 낸다(토큰 탐지용).

    토큰은 본문뿐 아니라 머리말/꼬리말·중첩표에도 들어갈 수 있다. 한 곳이라도
    놓치면 `has_tokens`가 False로 오판돼, 이미 토큰이 있는 양식이 자동 매핑 경로로
    잘못 흘러가 토큰이 이중 삽입된다. 그래서 탐지는 전 영역을 훑는다.
    """
    yield from _iter_container_texts(doc)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            yield from _iter_container_texts(hf)


def _has_tokens(doc) -> bool:
    """문서 전체(본문·머리말/꼬리말·중첩표)에 docxtpl 토큰(`{{`/`{%`)이 있으면 True."""
    return any("{{" in t or "{%" in t for t in _iter_all_texts(doc))


def _inspect_table(table, index: int) -> dict:
    """한 표의 셀 구조를 JSON 친화 dict로. 병합 원점만 출력.

    병합은 XML로 판정한다(lxml 프록시 id는 접근마다 달라 신뢰 못 함):
    - gridSpan > 1 → 가로 병합.
    - vMerge == 'restart' → 세로 병합 원점, 'continue' → 아래로 이어진 중복 셀(제외).
    셀의 논리 열은 tc를 왼쪽부터 훑으며 gridSpan만큼 누적해 계산한다.
    """
    from docx.table import _Cell

    n_rows = len(table.rows)
    n_cols = len(table.columns)

    cells = []
    for r, row in enumerate(table.rows):
        col = 0
        for tc in row._tr.tc_lst:
            span = tc.grid_span
            vmerge = tc.vMerge  # None | 'restart' | 'continue'
            if vmerge == "continue":
                col += span  # 세로 병합 중복 — 원점(위 restart 행)만 남긴다
                continue
            text = _Cell(tc, table).text.strip()
            cells.append({
                "row": r,
                "col": col,
                "text": text,
                "is_empty": text == "",
                "merged": span > 1 or vmerge == "restart",
            })
            col += span
    return {"index": index, "rows": n_rows, "cols": n_cols, "cells": cells}


def _inspect_paragraphs(doc) -> list:
    """본문 문단을 인덱스와 함께 덤프(문단 기반 양식 매핑용).

    인덱스는 doc.paragraphs 상의 위치이며 mapping.json의 `para` 주소와 일치한다.
    빈 문단(간격용)도 인덱스 정확성을 위해 그대로 포함한다. 표 안의 문단은 제외된다.
    """
    out = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        out.append({"index": i, "text": text, "is_empty": text == ""})
    return out


def inspect_template(template_path: str) -> dict:
    """.docx 양식을 열어 구조 JSON(dict)을 반환한다.

    표(`tables`)와 본문 문단(`paragraphs`)을 모두 덤프하므로, 표 기반·문단 기반
    양식 모두에 매핑을 만들 수 있다.
    """
    from docx import Document

    resolved = resolve_input_path(template_path, must_exist=True)
    if resolved.suffix.lower() != ".docx":
        raise ValueError(
            f"양식은 .docx여야 합니다: {resolved.name}. "
            "hwp/pdf 양식이면 한글/워드에서 .docx로 저장해 다시 주세요."
        )
    doc = Document(str(resolved))
    return {
        "has_tokens": _has_tokens(doc),
        "tables": [_inspect_table(t, i) for i, t in enumerate(doc.tables)],
        "paragraphs": _inspect_paragraphs(doc),
    }


def main() -> None:
    path = sys.argv[1]
    result = inspect_template(path)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
